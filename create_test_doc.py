#!/usr/bin/env python3
"""
Create a test DOCX document with intentional grammar errors for testing the API.
"""

from docx import Document
import os


def create_test_document():
	"""Create a test document with grammar errors."""
	doc = Document()
	
	# Add title
	title = doc.add_heading('Test Document for Grammar Check', 0)
	
	# Add paragraphs with intentional errors
	doc.add_paragraph('This are a test document.')
	doc.add_paragraph('I goes to the store every day.')
	doc.add_paragraph('The cat and dog is playing together.')
	doc.add_paragraph('She have three books.')
	doc.add_paragraph('They was at the party last night.')
	doc.add_paragraph('The weather look good today.')
	doc.add_paragraph('He don\'t like coffee.')
	doc.add_paragraph('The children is sleeping.')
	doc.add_paragraph('There is many people here.')
	doc.add_paragraph('The team are winning.')
	
	# Save the document
	output_path = "test_document.docx"
	doc.save(output_path)
	print(f"Test document created: {output_path}")
	print(f"File size: {os.path.getsize(output_path)} bytes")
	
	return output_path


if __name__ == "__main__":
	create_test_document()
