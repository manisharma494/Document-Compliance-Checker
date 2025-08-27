from __future__ import annotations
from io import BytesIO
from docx import Document
import language_tool_python
import re
from typing import List, Dict, Any


def clean_text_for_docx(text: str) -> str:
	"""Clean text to make it compatible with DOCX XML requirements."""
	if not text:
		return ""
	# Remove NULL bytes and control characters (except newlines and tabs)
	cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
	# Replace multiple newlines with single newlines
	cleaned = re.sub(r'\n\s*\n', '\n', cleaned)
	# Remove leading/trailing whitespace
	cleaned = cleaned.strip()
	# Ensure the text is not empty after cleaning
	if not cleaned:
		return "Document content"
	return cleaned


def generate_corrected_text(text: str, language: str = "en-US") -> str:
	"""Generate corrected text using LanguageTool."""
	try:
		tool = language_tool_python.LanguageToolPublicAPI(language)
		corrected = tool.correct(text)
		return corrected
	except Exception:
		return text  # Return original text if LanguageTool fails


def apply_ai_context_improvements(text: str, enhanced_issues: List[Dict[str, Any]]) -> str:
	"""Apply AI-generated context improvements to the text."""
	if not enhanced_issues:
		return text
	
	# Sort issues by offset in reverse order to avoid position shifting
	sorted_issues = sorted(enhanced_issues, key=lambda x: x.get('offset', 0), reverse=True)
	
	improved_text = text
	
	for issue in sorted_issues:
		if issue.get('context_improvement') and issue.get('enhanced'):
			offset = issue.get('offset', 0)
			length = issue.get('length', 0)
			improvement = issue['context_improvement']
			
			# Extract the problematic text
			problematic_text = improved_text[offset:offset + length]
			
			# Apply the AI improvement
			# The improvement might be a suggestion or a corrected version
			if improvement.startswith('"') and improvement.endswith('"'):
				# If improvement is quoted, extract the content
				replacement = improvement[1:-1]
			else:
				# Otherwise, use the improvement as-is
				replacement = improvement
			
			# Replace the problematic text with the improvement
			improved_text = improved_text[:offset] + replacement + improved_text[offset + length:]
	
	return improved_text


def generate_enhanced_corrected_text(text: str, enhanced_issues: List[Dict[str, Any]], language: str = "en-US") -> str:
	"""Generate enhanced corrected text with both grammar corrections and AI context improvements."""
	# First, apply basic grammar corrections
	grammar_corrected = generate_corrected_text(text, language)
	
	# Then, apply AI context improvements
	fully_enhanced = apply_ai_context_improvements(grammar_corrected, enhanced_issues)
	
	return fully_enhanced


def text_to_docx_bytes(text: str) -> bytes:
	"""Convert text to DOCX bytes."""
	doc = Document()
	# Clean the text first
	cleaned_text = clean_text_for_docx(text)
	# Split by newlines and add paragraphs
	paragraphs = cleaned_text.split('\n')
	for paragraph in paragraphs:
		# Skip empty paragraphs
		if paragraph.strip():
			doc.add_paragraph(paragraph.strip())
	# If no paragraphs were added, add a default one
	if len(doc.paragraphs) == 0:
		doc.add_paragraph("Document content")
	buffer = BytesIO()
	doc.save(buffer)
	return buffer.getvalue()


def create_enhanced_docx_with_annotations(text: str, enhanced_issues: List[Dict[str, Any]], language: str = "en-US") -> bytes:
	"""Create an enhanced DOCX with AI improvements and annotations."""
	doc = Document()
	
	# Add a title explaining the enhancements
	title = doc.add_heading('AI-Enhanced Document', 0)
	subtitle = doc.add_paragraph('This document has been enhanced with AI-powered context improvements and grammar corrections.')
	subtitle.alignment = 1  # Center alignment
	
	# Add a summary of improvements
	if enhanced_issues:
		summary_heading = doc.add_heading('Improvements Applied', level=1)
		summary_para = doc.add_paragraph('The following AI-powered improvements have been applied to enhance clarity and context:')
		
		for i, issue in enumerate(enhanced_issues, 1):
			if issue.get('enhanced') and issue.get('context_improvement'):
				issue_text = doc.add_paragraph(f'{i}. {issue.get("message", "Grammar issue")}')
				improvement_text = doc.add_paragraph(f'   AI Improvement: {issue["context_improvement"]}')
				improvement_text.style = 'Quote'
	
	# Add a separator
	doc.add_paragraph('---')
	
	# Apply corrections and improvements to the main text
	enhanced_text = generate_enhanced_corrected_text(text, enhanced_issues, language)
	
	# Add the enhanced content
	content_heading = doc.add_heading('Enhanced Content', level=1)
	
	# Split by newlines and add paragraphs
	paragraphs = enhanced_text.split('\n')
	for paragraph in paragraphs:
		if paragraph.strip():
			doc.add_paragraph(paragraph.strip())
	
	# If no content paragraphs were added, add a default one
	if len(doc.paragraphs) <= 3:  # Only title, subtitle, and headings
		doc.add_paragraph("Enhanced document content")
	
	buffer = BytesIO()
	doc.save(buffer)
	return buffer.getvalue()


def create_comparison_docx(original_text: str, enhanced_issues: List[Dict[str, Any]], language: str = "en-US") -> bytes:
	"""Create a comparison DOCX showing original vs. improved versions."""
	doc = Document()
	
	# Title
	title = doc.add_heading('Document Improvement Comparison', 0)
	subtitle = doc.add_paragraph('Original vs. AI-Enhanced Version')
	subtitle.alignment = 1
	
	# Original version
	doc.add_heading('Original Version', level=1)
	original_para = doc.add_paragraph(clean_text_for_docx(original_text))
	original_para.style = 'Quote'
	
	# Separator
	doc.add_paragraph('---')
	
	# Enhanced version
	doc.add_heading('AI-Enhanced Version', level=1)
	enhanced_text = generate_enhanced_corrected_text(original_text, enhanced_issues, language)
	enhanced_para = doc.add_paragraph(clean_text_for_docx(enhanced_text))
	enhanced_para.style = 'Quote'
	
	# Improvements summary
	if enhanced_issues:
		doc.add_heading('Applied Improvements', level=1)
		for i, issue in enumerate(enhanced_issues, 1):
			if issue.get('enhanced') and issue.get('context_improvement'):
				issue_para = doc.add_paragraph(f'{i}. Issue: {issue.get("message", "Grammar issue")}')
				improvement_para = doc.add_paragraph(f'   AI Enhancement: {issue["context_improvement"]}')
				improvement_para.style = 'Quote'
	
	buffer = BytesIO()
	doc.save(buffer)
	return buffer.getvalue()

